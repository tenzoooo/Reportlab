import sys
import os

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Trying to read xml directly.")
    sys.exit(1)

def extract_text_from_docx(docx_path):
    document = Document(docx_path)
    full_text = []
    
    # Paragraphs
    for para in document.paragraphs:
        full_text.append(para.text)
        
    # Tables
    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))
            
    return '\n'.join(full_text)

docx_path = "update_UI/report_backend/.agent_data/jobs/36606fa6a5874f8ea42fb5508c518001/artifact/report.docx"
if os.path.exists(docx_path):
    print(extract_text_from_docx(docx_path))
else:
    print(f"File not found: {docx_path}")
