
import sys
import os

# Add the directory containing the library to the python path
sys.path.append(os.path.join(os.getcwd(), 'lib/docx'))

from docxtpl import DocxTemplate
from docx.oxml.ns import qn
from render_with_docxtpl import build_table_subdoc

def verify_extended_units():
    template_path = 'templates/template.docx'
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = DocxTemplate(template_path)
    
    # Test cases:
    # 1. Existing bracket format: Vbe[V]
    # 2. Parentheses format: Length (m)
    # 3. Standalone format: kg
    # 4. Non-unit text: Description
    
    rows = [
        ["Vbe[V]", "Length (m)", "kg", "Description"],
        ["1", "2", "3", "4"]
    ]
    
    subdoc = build_table_subdoc(doc, rows)
    
    if not subdoc.tables:
        print("No tables found in subdoc")
        sys.exit(1)
        
    table = subdoc.tables[0]
    
    # Helper to check if cell has OMML
    def has_omml(cell):
        if not cell.paragraphs:
            return False
        p = cell.paragraphs[0]
        return p._p.find(qn("m:oMath")) is not None

    # Check Vbe[V] -> Should be OMML
    if not has_omml(table.cell(0, 0)):
        print("FAIL: Vbe[V] not converted to OMML")
        sys.exit(1)
        
    # Check Length (m) -> Should be OMML
    if not has_omml(table.cell(0, 1)):
        print("FAIL: Length (m) not converted to OMML")
        sys.exit(1)
        
    # Check kg -> Should be OMML
    if not has_omml(table.cell(0, 2)):
        print("FAIL: kg not converted to OMML")
        sys.exit(1)
        
    # Check Description -> Should NOT be OMML
    if has_omml(table.cell(0, 3)):
        print("FAIL: Description converted to OMML incorrectly")
        sys.exit(1)
            
    print("SUCCESS: Extended unit detection works correctly.")

if __name__ == "__main__":
    verify_extended_units()
