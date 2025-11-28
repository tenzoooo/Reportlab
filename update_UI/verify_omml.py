
import sys
import os

# Add the directory containing the library to the python path
sys.path.append(os.path.join(os.getcwd(), 'lib/docx'))

from docxtpl import DocxTemplate
from docx.oxml.ns import qn
from render_with_docxtpl import build_table_subdoc

def verify_omml_conversion():
    template_path = 'templates/template.docx'
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = DocxTemplate(template_path)
    
    rows = [
        ["Header 1", "Vbe[V]"],
        ["Cell 1", "Cell 2"]
    ]
    
    subdoc = build_table_subdoc(doc, rows)
    
    if not subdoc.tables:
        print("No tables found in subdoc")
        sys.exit(1)
        
    table = subdoc.tables[0]
    
    # Check if Vbe[V] cell has OMML
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    
    # Check for m:oMath element
    oMath = p._p.find(qn("m:oMath"))
    if oMath is None:
        print("FAIL: No OMML element found in unit cell")
        sys.exit(1)
        
    # Check text content inside OMML
    t = oMath.find(".//" + qn("m:t"))
    if t is None:
        print("FAIL: No text element found in OMML")
        sys.exit(1)
        
    if t.text != "Vbe[V]":
        print(f"FAIL: OMML text mismatch. Expected 'Vbe[V]', got '{t.text}'")
        sys.exit(1)
            
    print("SUCCESS: Unit text converted to OMML.")

if __name__ == "__main__":
    verify_omml_conversion()
