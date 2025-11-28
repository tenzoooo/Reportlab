import sys
from docx import Document

def extract_structure(docx_path):
    try:
        doc = Document(docx_path)
        print(f"--- Structure of {docx_path} ---")
        for para in doc.paragraphs:
            if para.text.strip():
                # Print style and text to understand hierarchy
                print(f"[{para.style.name}] {para.text}")
        
        print("\n--- Tables ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
            # Print first few cells to identify content
            if table.rows:
                first_row = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"  Header: {first_row}")

    except Exception as e:
        print(f"Error reading DOCX: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 extract_docx_structure.py <docx_path>")
        sys.exit(1)
    
    extract_structure(sys.argv[1])
