
import sys
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_omml_element(text):
    # Create OMML structure
    # <m:oMath>
    #   <m:r>
    #     <m:t>text</m:t>
    #   </m:r>
    # </m:oMath>
    
    oMath = OxmlElement('m:oMath')
    
    # Run
    r = OxmlElement('m:r')
    
    # Text
    t = OxmlElement('m:t')
    t.text = text
    
    r.append(t)
    oMath.append(r)
    
    return oMath

def test_omml_insertion():
    doc = Document()
    
    # Add a table
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add a paragraph to the cell
    p = cell.paragraphs[0]
    
    # Create OMML element
    omml = create_omml_element("Vbe[V]")
    
    # Append OMML to paragraph element
    # We need to access the internal xml element of the paragraph
    p._p.append(omml)
    
    output_path = 'test_omml.docx'
    doc.save(output_path)
    print(f"Saved {output_path}")

if __name__ == "__main__":
    test_omml_insertion()
