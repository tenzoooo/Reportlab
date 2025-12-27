from docx import Document

from templating.renderer import _normalize_jinja_tags


def test_normalize_merges_split_mustache_tag():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("prefix ")
    p.add_run("{{ chap")
    p.add_run("ter")
    p.add_run(" }} suffix")

    _normalize_jinja_tags(doc)

    joined = "".join(r.text or "" for r in p.runs)
    assert "{{ chapter }}" in joined
    assert any("{{ chapter }}" in (r.text or "") for r in p.runs)


def test_normalize_merges_split_block_tag():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{% fo")
    p.add_run("r x in xs ")
    p.add_run("%}")

    _normalize_jinja_tags(doc)

    joined = "".join(r.text or "" for r in p.runs)
    assert "{% for x in xs %}" in joined
    assert any("{% for x in xs %}" in (r.text or "") for r in p.runs)

