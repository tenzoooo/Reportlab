from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
import subprocess
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

from pydantic import BaseModel, Field

from core.past_report import extract_past_report_text, ext_from_filename
from core.storage import Storage
from graph.state import AgentState, PastReportData, now_iso
from llm.client import LLMClient
from llm.schemas.past_report_result_headings_select import PastReportResultHeadingsSelectOutput
from docx import Document
from docx.oxml.ns import qn
import fitz  # PyMuPDF
from models.contracts import (
    PastReportExperimentHint,
    PastReportGraphHint,
    PastReportHeading,
    PastReportTableColumnHint,
    PastReportTableHint,
)

_ERR_MISSING_STORAGE = "past_report_missing_storage"
_ERR_READ_FAILED = "past_report_read_failed"
_ERR_LLM_MISSING = "past_report_llm_missing"
_ERR_LLM_FAILED = "past_report_llm_failed"
_ERR_TEXT_EMPTY = "past_report_text_empty"
_DEFAULT_PARALLEL_WORKERS = 2
_DOCX_PDF_TIMEOUT_SEC = 120
_PAGE_IMAGE_ZOOM = 2.0
_MAX_PAGES_PER_SECTION = 30

logger = logging.getLogger(__name__)

def _ensure_past_reports(state: AgentState) -> list[PastReportData]:
    if state.past_reports:
        return list(state.past_reports)
    if state.past_report.storage_key:
        if not state.past_report.report_id:
            state.past_report.report_id = "legacy"
        state.past_reports = [state.past_report]
        return list(state.past_reports)
    return []

def _set_report_error(report: PastReportData, code: str) -> None:
    report.hints_error = code
    report.hints_ready = True


class _PastReportHintsOutput(BaseModel):
    hints: list[PastReportExperimentHint] = Field(default_factory=list)


class _PastReportHeadingsOutput(BaseModel):
    items: list[PastReportHeading] = Field(default_factory=list)


def _build_past_report_headings_messages(
    *,
    past_report_text: str,
    experiments: list[dict[str, str]],
) -> list[dict]:
    system = (
        "あなたは過去レポート本文から「見出し行」を抽出する抽出器です。\n"
        "出力はJSONのみ（説明文は禁止）。\n\n"
        "# ルール\n"
        "- heading_line は必ず本文中の部分文字列（完全一致）。\n"
        "- section_number は必ず設定する（必須）。\n"
        "- section_number はドット区切り（例: 4.3.1 / 1.2）。ハイフン等はドットに正規化する。\n"
        "- 本文に番号がない見出しは、前後の文脈から推定して section_number を付与する。\n"
        "- level は 1-6 の整数。番号付きは階層に対応。\n"
        "- title は見出しの本文タイトル。\n"
        "- 図表番号やページ番号、ヘッダ/フッタは除外。\n\n"
        "# 出力\n"
        "{\n"
        "  \"items\": [\n"
        "    {\"heading_line\": \"...\", \"level\": 1, \"section_number\": \"\", \"title\": \"...\"}\n"
        "  ]\n"
        "}\n"
    )
    return [
        {"role": "system", "content": system},
        {
            "role": "user",
            "content": json.dumps(
                {"past_report_text": past_report_text, "experiments": experiments},
                ensure_ascii=False,
            ),
        },
    ]


def _build_result_heading_select_messages(*, heading_lines: list[str], experiments: list[dict[str, str]]) -> list[dict]:
    system = (
        "あなたは見出し一覧から「実験結果＋結果に準ずる節」を選別する抽出器です。\n"
        "出力はJSONのみ（説明文は禁止）。\n\n"
        "# 入力\n"
        "- heading_lines: 見出し行の配列\n\n"
        "- experiments: Bレイヤの実験リスト（exp_key, title）\n\n"
        "# 出力\n"
        "{\n"
        "  \"heading_lines\": [\"...\"]\n"
        "}\n\n"
        "# ルール\n"
        "- 実験結果または結果に準ずる節のみ選ぶ。\n"
        "- 見出し行は入力配列の要素から完全一致で選ぶ。\n"
        "- 順序は入力の順序を保持する。\n"
        "- 見出しが無い場合は空配列。\n"
    )
    return [
        {"role": "system", "content": system},
        {
            "role": "user",
            "content": json.dumps(
                {"heading_lines": heading_lines, "experiments": experiments},
                ensure_ascii=False,
            ),
        },
    ]


def _normalize_section_number(raw: str) -> str:
    s = (raw or "").strip().replace("．", ".")
    s = s.translate(str.maketrans("０１２３４５６７８９", "0123456789"))
    s = re.sub(r"[\\-‐‑–—−ー]+", ".", s)
    s = re.sub(r"\\.+", ".", s)
    return s.strip(".")


def _extract_section_number_from_heading_line(text: str) -> str:
    line = (text or "").strip().replace("．", ".")
    line = line.translate(str.maketrans("０１２３４５６７８９", "0123456789"))
    match = re.match(r"^\\s*(?P<section>\\d+(?:[.\\-‐‑–—−ー]\\d+)+|\\d+)\\s*(?P<title>.+)?$", line)
    if not match:
        return ""
    return _normalize_section_number(match.group("section") or "")


def _split_heading_line(text: str) -> tuple[str, str]:
    line = (text or "").strip()
    if not line:
        return "", ""
    match = re.match(r"^\\s*(?P<section>\\d+(?:[.\\-‐‑–—−ー]\\d+)+|\\d+)\\s*(?P<title>.+)$", line)
    if not match:
        return "", line
    section = _normalize_section_number(match.group("section") or "")
    title = (match.group("title") or "").strip()
    return section, title


def _infer_exp_key(*, section_number: str, heading_line: str) -> str:
    if section_number:
        return section_number
    return _extract_section_number_from_heading_line(heading_line)


def _normalize_headings(items: list[PastReportHeading]) -> list[PastReportHeading]:
    normalized: list[PastReportHeading] = []
    for item in items:
        heading_line = (item.heading_line or "").strip()
        section_number = _normalize_section_number(item.section_number or "")
        title = (item.title or "").strip()
        if not section_number or not title:
            extracted_section, extracted_title = _split_heading_line(heading_line)
            if not section_number and extracted_section:
                section_number = extracted_section
            if not title and extracted_title:
                title = extracted_title
        if not section_number or not title or not heading_line:
            continue
        level = item.level if 1 <= (item.level or 0) <= 6 else max(1, min(6, section_number.count(".") + 1))
        exp_key = _infer_exp_key(section_number=section_number, heading_line=heading_line)
        normalized.append(
            PastReportHeading(
                heading_line=heading_line,
                level=level,
                section_number=section_number,
                exp_key=exp_key,
                title=title,
                section_text=item.section_text,
            )
        )
    return normalized


def _build_past_report_hints_messages(*, past_report_text: str) -> list[dict]:
    system = (
        "あなたは過去レポート本文から、各実験の表・グラフ情報を抽出する抽出器です。\n"
        "出力はJSONのみ（説明文は禁止）。\n\n"
        "# 入力\n"
        "- past_report_text: 過去レポート本文\n\n"
        "# 出力仕様\n"
        "{\n"
        "  \"hints\": [\n"
        "    {\n"
        "      \"exp_key\": \"4.2.1\",\n"
        "      \"name\": \"反転増幅回路\",\n"
        "      \"tables_count\": 2,\n"
        "      \"graphs_count\": 1,\n"
        "      \"tables\": [\n"
        "        {\"caption\": \"...\", \"columns\": [{\"name\": \"\", \"unit\": \"\"}]}\n"
        "      ],\n"
        "      \"graphs\": [\n"
        "        {\"x_name\": \"\", \"x_unit\": \"\", \"y_name\": \"\", \"y_unit\": \"\", \"series_names\": [], \"condition_names\": [], \"caption\": \"\"}\n"
        "      ]\n"
        "    }\n"
        "  ]\n"
        "}\n\n"
        "# ルール\n"
        "- 本文に現れる実験を対象にする。\n"
        "- 不明な項目は空文字や空配列でよい。\n"
        "- tables_count/graphs_count は tables/graphs の件数と一致させる。\n"
        "- 表/図のキャプションが本文にあれば可能な限り抽出する。\n"
    )
    return [
        {"role": "system", "content": system},
        {
            "role": "user",
            "content": json.dumps({"past_report_text": past_report_text}, ensure_ascii=False),
        },
    ]


def _normalize_hints(hints: list[PastReportExperimentHint]) -> list[PastReportExperimentHint]:
    normalized: list[PastReportExperimentHint] = []
    for hint in hints:
        graphs = []
        tables = []
        for g in hint.graphs or []:
            graphs.append(
                PastReportGraphHint(
                    x_name=g.x_name or "",
                    x_unit=g.x_unit or "",
                    y_name=g.y_name or "",
                    y_unit=g.y_unit or "",
                    series_names=list(g.series_names or []),
                    condition_names=list(g.condition_names or []),
                    caption=g.caption or "",
                )
            )
        for t in hint.tables or []:
            cols = []
            for c in t.columns or []:
                cols.append(PastReportTableColumnHint(name=c.name or "", unit=c.unit or ""))
            tables.append(PastReportTableHint(columns=cols, caption=t.caption or ""))
        normalized.append(
            PastReportExperimentHint(
                exp_key=hint.exp_key or "",
                name=hint.name or "",
                graphs_count=max(0, int(hint.graphs_count or 0)),
                tables_count=max(0, int(hint.tables_count or 0)),
                graphs=graphs,
                tables=tables,
            )
        )
    return normalized


def _collect_experiment_list(state: AgentState) -> list[dict[str, str]]:
    experiments: list[dict[str, str]] = []
    seen: set[str] = set()
    bundle = state.b_layer_bundle
    method_items = (bundle.method.items if bundle and bundle.method else []) or []
    for item in method_items:
        key = str(item.exp_key or "").strip()
        title = str(item.title or "").strip()
        if not key or key in seen:
            continue
        seen.add(key)
        experiments.append({"exp_key": key, "title": title})
    if experiments:
        return experiments
    if not bundle or not bundle.theory_formulas:
        return experiments
    for item in bundle.theory_formulas:
        for exp in item.experiments or []:
            key = str(exp or "").strip()
            if not key or key in seen:
                continue
            seen.add(key)
            experiments.append({"exp_key": key, "title": ""})
    return experiments


def _llm_parallel_workers() -> int:
    raw = (os.environ.get("REPORT_AGENT_LLM_PARALLEL_WORKERS") or "").strip()
    if raw.isdigit():
        return max(1, int(raw))
    return _DEFAULT_PARALLEL_WORKERS


def _run_hint_for_section(
    *,
    llm: LLMClient,
    section_text: str,
    image_data_urls: list[str],
) -> list[PastReportExperimentHint]:
    hints_output = llm.parse(
        _PastReportHintsOutput,
        model=llm.text_model,
        messages=_build_past_report_hints_messages(past_report_text=section_text),
        attempts=2,
    )
    return _normalize_hints(hints_output.hints or [])


def _find_heading_start(text: str, *, heading_line: str, title: str, cursor: int) -> int | None:
    if not text:
        return None
    if heading_line:
        idx = text.find(heading_line, cursor)
        if idx >= 0:
            return idx
    if title:
        pattern = re.compile(rf"(?m)^\\s*.*{re.escape(title)}.*$")
        match = pattern.search(text, cursor)
        if match:
            return match.start()
    return None


def _attach_section_text(text: str, headings: list[PastReportHeading]) -> list[PastReportHeading]:
    if not text or not headings:
        return headings
    out: list[PastReportHeading] = []
    cursor = 0
    for idx, item in enumerate(headings):
        heading_line = item.heading_line or ""
        title = item.title or ""
        start_idx = _find_heading_start(text, heading_line=heading_line, title=title, cursor=cursor)
        if start_idx is None:
            out.append(item)
            continue
        section_start = start_idx + len(heading_line)
        next_idx = len(text)
        if idx + 1 < len(headings):
            next_heading = headings[idx + 1].heading_line or ""
            if next_heading:
                found = text.find(next_heading, section_start)
                if found >= 0:
                    next_idx = found
        section_text = text[section_start:next_idx].strip()
        out.append(
            PastReportHeading(
                heading_line=item.heading_line,
                level=item.level,
                section_number=item.section_number,
                exp_key=item.exp_key or item.section_number,
                title=item.title,
                section_text=section_text,
            )
        )
        cursor = section_start
    return out


def _safe_filename(value: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9._-]+", "_", value or "")
    return s.strip("_") or "section"


def _image_ext_from_content_type(content_type: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
        "image/webp": ".webp",
    }
    return mapping.get(content_type, ".bin")


def _attach_docx_images_to_result_headings(
    *,
    report: PastReportData,
    raw_docx: bytes,
    base_dir: Path,
) -> None:
    doc = Document(io.BytesIO(raw_docx))
    paragraphs = list(doc.paragraphs)
    para_texts = [p.text or "" for p in paragraphs]
    para_images: list[list[str]] = []
    for p in paragraphs:
        rids: list[str] = []
        for blip in p._element.xpath(".//*[local-name()='blip']"):
            rid = blip.get(qn("r:embed"))
            if rid:
                rids.append(rid)
        para_images.append(rids)

    heading_indices: list[int | None] = []
    cursor = 0
    for heading in report.result_section_headings:
        line = (heading.heading_line or "").strip()
        if not line:
            heading_indices.append(None)
            continue
        found = None
        for idx in range(cursor, len(para_texts)):
            if line in para_texts[idx]:
                found = idx
                cursor = idx + 1
                break
        heading_indices.append(found)

    base_dir.mkdir(parents=True, exist_ok=True)
    for idx, heading in enumerate(report.result_section_headings):
        start = heading_indices[idx]
        if start is None:
            continue
        end = None
        for next_idx in heading_indices[idx + 1 :]:
            if next_idx is not None:
                end = next_idx
                break
        if end is None:
            end = len(para_images)
        rids: list[str] = []
        for j in range(start, end):
            for rid in para_images[j]:
                if rid not in rids:
                    rids.append(rid)
        if not rids:
            continue
        image_paths: list[str] = []
        image_data_urls: list[str] = []
        for rid_idx, rid in enumerate(rids, start=1):
            part = doc.part.related_parts.get(rid)
            if part is None:
                continue
            content_type = getattr(part, "content_type", "") or ""
            ext = _image_ext_from_content_type(content_type)
            section_id = heading.section_number or f"{idx+1}"
            filename = f"{_safe_filename(section_id)}_{rid_idx}{ext}"
            path = base_dir / filename
            blob = part.blob
            path.write_bytes(blob)
            image_paths.append(str(path))
            if content_type:
                b64 = base64.b64encode(blob).decode("ascii")
                image_data_urls.append(f"data:{content_type};base64,{b64}")
        heading.image_paths = image_paths
        heading.image_data_urls = image_data_urls


def _convert_docx_to_pdf(*, docx_path: Path, out_dir: Path) -> Path | None:
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        "/opt/homebrew/bin/soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(docx_path),
    ]
    try:
        subprocess.run(cmd, check=True, timeout=_DOCX_PDF_TIMEOUT_SEC)
    except Exception as exc:
        logger.warning("past_report_hints: docx->pdf failed: %s", exc)
        return None
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    return pdf_path if pdf_path.exists() else None


def _find_heading_pages(pdf_path: Path, headings: list[PastReportHeading]) -> list[int | None]:
    doc = fitz.open(pdf_path)
    pages_text = [page.get_text("text") or "" for page in doc]
    indices: list[int | None] = []
    cursor = 0
    for heading in headings:
        line = (heading.heading_line or "").strip()
        title = (heading.title or "").strip()
        found = None
        for idx in range(cursor, len(pages_text)):
            text = pages_text[idx]
            if (line and line in text) or (title and title in text):
                found = idx
                cursor = idx + 1
                break
        indices.append(found)
    return indices


def _attach_page_images_to_result_headings(
    *,
    report: PastReportData,
    pdf_path: Path,
    base_dir: Path,
) -> None:
    base_dir.mkdir(parents=True, exist_ok=True)
    headings = list(report.result_section_headings or [])
    if not headings:
        return
    page_indices = _find_heading_pages(pdf_path, headings)
    doc = fitz.open(pdf_path)
    for idx, heading in enumerate(headings):
        start = page_indices[idx]
        if start is None:
            continue
        end = None
        for next_idx in page_indices[idx + 1 :]:
            if next_idx is not None:
                end = next_idx - 1
                break
        if end is None:
            end = doc.page_count - 1
        if end < start:
            continue
        page_image_paths: list[str] = []
        count = 0
        for page_no in range(start, end + 1):
            if count >= _MAX_PAGES_PER_SECTION:
                break
            page = doc.load_page(page_no)
            pix = page.get_pixmap(matrix=fitz.Matrix(_PAGE_IMAGE_ZOOM, _PAGE_IMAGE_ZOOM), alpha=False)
            png_bytes = pix.tobytes("png")
            section_id = heading.section_number or f"{idx+1}"
            filename = f"{_safe_filename(section_id)}_p{page_no+1}.png"
            path = base_dir / filename
            path.write_bytes(png_bytes)
            page_image_paths.append(str(path))
            count += 1
        heading.page_image_paths = page_image_paths
        heading.page_image_data_urls = []


def _extract_result_section_headings(headings: list[PastReportHeading]) -> list[PastReportHeading]:
    if not headings:
        return []
    start_idx = None
    end_idx = None
    for idx, h in enumerate(headings):
        title = (h.title or h.heading_line or "")
        if start_idx is None and ("実験結果" in title or "結果" in title):
            start_idx = idx
            continue
        if start_idx is not None and ("考察" in title or "検討" in title):
            end_idx = idx
            break
    if start_idx is None:
        return []
    if end_idx is None:
        end_idx = len(headings)
    return headings[start_idx:end_idx]


def _select_result_section_headings_with_llm(
    *,
    headings: list[PastReportHeading],
    experiments: list[dict[str, str]],
    llm: LLMClient,
) -> list[PastReportHeading]:
    if not headings:
        return []
    lines = [str(h.heading_line or "").strip() for h in headings if str(h.heading_line or "").strip()]
    if not lines:
        return []
    result = llm.parse(
        PastReportResultHeadingsSelectOutput,
        model="gpt-5-mini",
        messages=_build_result_heading_select_messages(heading_lines=lines, experiments=experiments),
        attempts=2,
    )
    chosen = [str(x or "").strip() for x in (result.heading_lines or []) if str(x or "").strip()]
    if not chosen:
        return []
    chosen_set = set(chosen)
    return [h for h in headings if (h.heading_line or "").strip() in chosen_set]

def past_report_hints(state: AgentState, *, storage: Storage, llm: LLMClient | None) -> AgentState:
    reports = _ensure_past_reports(state)
    if not reports:
        logger.info("past_report_hints: no past reports to process")
        return state

    changed = False
    for report in reports:
        if report.hints_ready:
            logger.info("past_report_hints: report %s already ready, skip", report.report_id or report.storage_key)
            continue
        if not report.storage_key:
            _set_report_error(report, _ERR_MISSING_STORAGE)
            changed = True
            continue
        try:
            raw = storage.get_bytes(report.storage_key)
        except Exception:
            _set_report_error(report, _ERR_READ_FAILED)
            changed = True
            continue

        ext = ext_from_filename(report.filename or report.storage_key)
        if ext not in {".pdf", ".docx"}:
            ext = ".pdf"
        text = extract_past_report_text(raw, ext=ext)
        if not text:
            _set_report_error(report, _ERR_TEXT_EMPTY)
            changed = True
            continue
        if llm is None:
            _set_report_error(report, _ERR_LLM_MISSING)
            changed = True
            continue

        experiments = _collect_experiment_list(state)
        headings_output = llm.parse(
            _PastReportHeadingsOutput,
            model="gpt-5-mini",
            messages=_build_past_report_headings_messages(
                past_report_text=text,
                experiments=experiments,
            ),
            attempts=2,
        )
        report.headings = _attach_section_text(text, _normalize_headings(list(headings_output.items or [])))
        selected = _select_result_section_headings_with_llm(
            headings=report.headings,
            experiments=experiments,
            llm=llm,
        )
        report.result_section_headings = selected or _extract_result_section_headings(report.headings)
        if ext == ".docx" and report.result_section_headings:
            report_id = report.report_id or Path(report.filename).stem or "past_report"
            base_root = Path(__file__).resolve().parents[2] / ".agent_data" / "past_report_images" / _safe_filename(report_id)
            base_root.mkdir(parents=True, exist_ok=True)
            docx_path = base_root / "source.docx"
            if not docx_path.exists():
                docx_path.write_bytes(raw)
            # Embedded images are intentionally skipped; we rely on page images only.
            pdf_path = base_root / "source.pdf"
            if not pdf_path.exists():
                converted = _convert_docx_to_pdf(docx_path=docx_path, out_dir=base_root)
                if converted is not None:
                    pdf_path = converted
            if pdf_path.exists():
                _attach_page_images_to_result_headings(
                    report=report,
                    pdf_path=pdf_path,
                    base_dir=base_root / "pages",
                )

        hints: list[PastReportExperimentHint] = []
        section_headings = list(report.result_section_headings or [])
        if section_headings:
            logger.info(
                "past_report_hints: LLM parse start report=%s sections=%s",
                report.report_id or report.storage_key,
                len(section_headings),
            )
            sections: list[tuple[int, PastReportHeading, str, list[str]]] = []
            for idx, heading in enumerate(section_headings):
                section_text = str(heading.section_text or "").strip()
                if not section_text:
                    continue
                image_data_urls = list(heading.page_image_data_urls or [])
                sections.append((idx, heading, section_text, image_data_urls))

            results: dict[int, PastReportExperimentHint] = {}
            failed = False
            workers = min(_llm_parallel_workers(), max(1, len(sections)))
            try:
                with ThreadPoolExecutor(max_workers=workers) as executor:
                    futures = {
                        executor.submit(
                            _run_hint_for_section,
                            llm=llm,
                            section_text=section_text,
                            image_data_urls=image_data_urls,
                        ): idx
                        for idx, _, section_text, image_data_urls in sections
                    }
                    for future in as_completed(futures):
                        idx = futures[future]
                        try:
                            chunk_hints = future.result()
                        except Exception as exc:
                            logger.warning("past_report_hints: LLM failed: %s", exc)
                            failed = True
                            continue
                        if not chunk_hints:
                            continue
                        results[idx] = chunk_hints[0]
            except Exception as exc:
                logger.warning("past_report_hints: LLM failed: %s", exc)
                failed = True

            for idx, heading, _, _ in sections:
                picked = results.get(idx)
                if picked is None:
                    continue
                if not picked.exp_key:
                    picked.exp_key = str(heading.section_number or "").strip()
                if not picked.name:
                    picked.name = str(heading.title or "").strip()
                hints.append(picked)
            if failed and not report.hints_error:
                report.hints_error = _ERR_LLM_FAILED
        else:
            logger.info(
                "past_report_hints: LLM parse start report=%s text_len=%s",
                report.report_id or report.storage_key,
                len(text),
            )
            hints_output = llm.parse(
                _PastReportHintsOutput,
                model=llm.text_model,
                messages=_build_past_report_hints_messages(past_report_text=text),
                attempts=2,
            )
            hints = _normalize_hints(hints_output.hints or [])

        report.hints = _normalize_hints(hints)
        logger.info(
            "past_report_hints: LLM parse done report=%s hints=%s",
            report.report_id or report.storage_key,
            len(report.hints),
        )
        report.hints_error = ""
        report.hints_ready = True
        changed = True

    if changed:
        state.job_meta.updated_at = now_iso()
    return state


def past_report_hints_and_save(
    state: AgentState,
    *,
    storage: Storage,
    llm: LLMClient | None,
    output_state_path: Path,
) -> AgentState:
    """
    Run past_report_hints and persist updated state to output_state_path.
    """
    state = past_report_hints(state, storage=storage, llm=llm)
    output_state_path.write_text(
        json.dumps(state.model_dump(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return state
