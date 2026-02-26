from __future__ import annotations

import re
from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Mm
from docx.shared import RGBColor

from core.storage import Storage
from models.contracts import ImageAsset


TARGET_WIDTH_MM = 106.29
TARGET_HEIGHT_MM = 60.57
_MISSING_CELL_MARK = "＼"

_IMG_RE = re.compile(r"^!\[(?P<alt>.*)\]\((?P<src>[^)]+)\)\s*$")
_HEADING_RE = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.+?)\s*$")
_CENTER_P_RE = re.compile(r"^<p\\s+align=[\"']center[\"']\\s*>(?P<text>.*?)</p>\\s*$", re.IGNORECASE)
_CENTER_DIV_RE = re.compile(r"^<div\\s+align=[\"']center[\"']\\s*>(?P<text>.*?)</div>\\s*$", re.IGNORECASE)
_CENTER_P_ANY_RE = re.compile(r"<p\\s+align=[\"']center[\"']\\s*>(?P<text>.*?)</p>", re.IGNORECASE)
_CENTER_DIV_ANY_RE = re.compile(r"<div\\s+align=[\"']center[\"']\\s*>(?P<text>.*?)</div>", re.IGNORECASE)
_CAPTIONISH_RE = re.compile(r"^(図|表)\\s*\\d+(?:\\.\\d+){0,3}(?:\\.\\d+)?\\s*[:：]?\\s*.+")


def _is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    buf: list[str] = []
    i = start
    while i < len(lines) and _is_table_line(lines[i]):
        buf.append(lines[i].strip())
        i += 1

    rows: list[list[str]] = []
    for raw in buf:
        # Remove leading/trailing | then split.
        parts = [p.strip().replace("\\|", "|") for p in raw.strip()[1:-1].split("|")]
        rows.append(parts)

    # Drop separator row like |---|---|
    if len(rows) >= 2 and all(set((c or "").replace(":", "").replace("-", "").replace(" ", "")) <= {""} for c in rows[1]):
        rows.pop(1)

    return rows, i


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max((len(r) for r in rows), default=0)
    if cols <= 0:
        return
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        row_list = row or []
        for c_idx in range(cols):
            val = ""
            if c_idx < len(row_list):
                val = "" if row_list[c_idx] is None else str(row_list[c_idx])
            cell = table.cell(r_idx, c_idx)
            if _is_missing_cell_value(val):
                cell.text = ""
                _apply_missing_cell_diagonal(cell)
            else:
                cell.text = val


def _is_missing_cell_value(value: str) -> bool:
    s = (value or "").strip()
    return not s or s == _MISSING_CELL_MARK


def _apply_missing_cell_diagonal(cell) -> None:
    # Draw a top-right to bottom-left diagonal in empty cells.
    try:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        diag = borders.find(qn("w:tr2bl"))
        if diag is None:
            diag = OxmlElement("w:tr2bl")
            borders.append(diag)
        diag.set(qn("w:val"), "single")
        diag.set(qn("w:sz"), "4")
        diag.set(qn("w:space"), "0")
        diag.set(qn("w:color"), "auto")
    except Exception:
        return


def _apply_mincho_black(doc: Document, *, font_name: str = "游明朝") -> None:
    """
    Force all text in the generated DOCX to:
    - black color
    - Mincho font (Japanese serif)
    This is applied as a post-pass over runs and also via styles as best-effort.
    """
    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        normal.font.color.rgb = RGBColor(0, 0, 0)
        # Set East Asian font for Normal.
        rfonts = normal._element.rPr.rFonts  # type: ignore[attr-defined]
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title", "Subtitle"]:
        try:
            st = doc.styles[style_name]
            st.font.name = font_name
            st.font.color.rgb = RGBColor(0, 0, 0)
            rfonts = st._element.rPr.rFonts  # type: ignore[attr-defined]
            rfonts.set(qn("w:eastAsia"), font_name)
        except Exception:
            continue

    def set_run(run) -> None:
        try:
            run.font.name = font_name
            run.font.color.rgb = RGBColor(0, 0, 0)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), font_name)
        except Exception:
            return

    try:
        for p in doc.paragraphs:
            for r in p.runs:
                set_run(r)
    except Exception:
        pass

    try:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            set_run(r)
    except Exception:
        pass


def _resolve_image_bytes(
    src: str,
    *,
    images_by_id: dict[str, ImageAsset],
    storage: Storage,
) -> bytes | None:
    s = (src or "").strip()
    if s.startswith("image:"):
        image_id = s[len("image:") :].strip()
        asset = images_by_id.get(image_id)
        if not asset:
            return None
        return storage.get_bytes(asset.storage_key)
    return None


def _normalize_omml(omml: str) -> str:
    s = (omml or "").strip()
    if not s:
        return ""
    if "xmlns:m" in s:
        return s
    if s.startswith("<m:oMath"):
        return s.replace("<m:oMath", '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"', 1)
    return s


def _append_omml(paragraph, omml: str) -> None:
    cleaned = _normalize_omml(omml)
    if not cleaned:
        return
    try:
        omml_xml = parse_xml(cleaned)
        paragraph._p.append(omml_xml)
    except Exception:
        return


def render_docx_from_markdown(
    *,
    markdown: str,
    storage: Storage,
    images_by_id: dict[str, ImageAsset] | None = None,
    omml_by_text: dict[str, list[str]] | None = None,
) -> bytes:
    """
    Minimal markdown → DOCX renderer for our agent-generated markdown.
    Supports:
    - Headings (#..######)
    - Paragraphs
    - Markdown tables (pipe format)
    - Images in the form: ![alt](image:<image_id>)
    """
    images_by_id = images_by_id or {}
    omml_by_text = omml_by_text or {}

    doc = Document()
    lines = (markdown or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        # Normalize inline HTML captions before any other parsing.
        if "<p" in line or "<div" in line:
            mcenter = _CENTER_P_ANY_RE.search(line) or _CENTER_DIV_ANY_RE.search(line)
            if mcenter:
                text = (mcenter.group("text") or "").strip()
                p = doc.add_paragraph(text)
                try:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
                i += 1
                continue
            # Strip unknown/unsupported HTML tags to avoid raw tags in output.
            line = re.sub(r"</?p[^>]*>", "", line, flags=re.IGNORECASE)
            line = re.sub(r"</?div[^>]*>", "", line, flags=re.IGNORECASE)
        if not line.strip():
            i += 1
            continue

        # Caption-like plain line (e.g., 図5.2.1.2 ... / 表5.2.1.1 ...)
        if _CAPTIONISH_RE.match(line.strip()):
            p = doc.add_paragraph(line.strip())
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            i += 1
            continue

        # Centered paragraph (HTML-ish), used for captions.
        stripped = line.strip()
        mcenter = None
        if stripped.lower().startswith("<p align=\"center\">") and stripped.lower().endswith("</p>"):
            inner = stripped[len("<p align=\"center\">") : -len("</p>")]
            mcenter = _CENTER_P_ANY_RE.search(f"<p align=\"center\">{inner}</p>")
        elif stripped.lower().startswith("<p align='center'>") and stripped.lower().endswith("</p>"):
            inner = stripped[len("<p align='center'>") : -len("</p>")]
            mcenter = _CENTER_P_ANY_RE.search(f"<p align=\"center\">{inner}</p>")
        elif stripped.lower().startswith("<div align=\"center\">") and stripped.lower().endswith("</div>"):
            inner = stripped[len("<div align=\"center\">") : -len("</div>")]
            mcenter = _CENTER_DIV_ANY_RE.search(f"<div align=\"center\">{inner}</div>")
        elif stripped.lower().startswith("<div align='center'>") and stripped.lower().endswith("</div>"):
            inner = stripped[len("<div align='center'>") : -len("</div>")]
            mcenter = _CENTER_DIV_ANY_RE.search(f"<div align=\"center\">{inner}</div>")
        if not mcenter:
            if "<p" in stripped or "<div" in stripped:
                if "<p" in stripped and "</p>" in stripped:
                    mcenter = _CENTER_P_ANY_RE.search(stripped)
                else:
                    mcenter = _CENTER_DIV_ANY_RE.search(stripped)
        if not mcenter:
            mcenter = _CENTER_P_RE.match(stripped) or _CENTER_DIV_RE.match(stripped)
        if mcenter:
            text = (mcenter.group("text") or "").strip()
            p = doc.add_paragraph(text)
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            i += 1
            continue

        # Headings
        m = _HEADING_RE.match(line)
        if m:
            level = min(6, len(m.group("hashes")))
            text = m.group("text").strip()
            # python-docx supports heading levels 0..9; use 1..4 for readability.
            doc.add_heading(text, level=min(4, level))
            i += 1
            continue

        # Tables
        if _is_table_line(line):
            rows, next_i = _parse_table_block(lines, i)
            _add_table(doc, rows)
            i = next_i
            continue

        # Images
        im = _IMG_RE.match(line.strip())
        if im:
            alt = (im.group("alt") or "").strip()
            src = (im.group("src") or "").strip()
            raw = _resolve_image_bytes(src, images_by_id=images_by_id, storage=storage)
            if raw:
                doc.add_picture(BytesIO(raw), width=Mm(TARGET_WIDTH_MM), height=Mm(TARGET_HEIGHT_MM))
                try:
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            else:
                doc.add_paragraph(f"[missing image: {src}]")
            i += 1
            continue

        # Paragraph block: join consecutive non-empty, non-special lines.
        buf: list[str] = [line.strip()]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip():
                break
            if _HEADING_RE.match(nxt) or _is_table_line(nxt) or _IMG_RE.match(nxt.strip()):
                break
            buf.append(nxt.strip())
            j += 1
        text = " ".join(buf).strip()
        p = doc.add_paragraph()
        omml_list = omml_by_text.get(text) or []
        if omml_list:
            omml = omml_list.pop(0)
            _append_omml(p, omml)
            if not omml_list:
                omml_by_text.pop(text, None)
        if text:
            p.add_run(text)
        if _CAPTIONISH_RE.match(text):
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        i = j + 1 if j < len(lines) and not lines[j].strip() else j

    _apply_mincho_black(doc, font_name="游明朝")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
