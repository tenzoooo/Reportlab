from __future__ import annotations

import logging
import os
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from docxtpl import DocxTemplate, InlineImage, RichText
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
from jinja2 import Environment, TemplateSyntaxError

from core.storage import Storage
from models.contracts import ImageAsset
from templating.filters import consideration_units, nl2br, reference_lines


logger = logging.getLogger(__name__)

TARGET_WIDTH_MM = 106.29
TARGET_HEIGHT_MM = 60.57
_MISSING_CELL_MARK = "＼"


def _build_env() -> Environment:
    env = Environment(autoescape=False)
    env.filters["nl2br"] = nl2br
    env.filters["consideration_units"] = consideration_units
    env.filters["reference_lines"] = reference_lines
    return env


def _create_consideration_units_rt(units: Any) -> RichText:
    return nl2br(consideration_units(units if isinstance(units, list) else []))


def _create_references_rt(consideration: Any) -> RichText:
    return nl2br(reference_lines(consideration if isinstance(consideration, dict) else {}))


def _patch_template(doc: DocxTemplate, context: dict) -> None:
    docx_obj = doc.get_docx()
    if not docx_obj:
        return

    _normalize_jinja_tags(docx_obj)

    replacements = {
        "{{ consideration.units | consideration_units | nl2br }}": "{{ consideration_units_rt }}",
        "{{ consideration.units | consideration_units }}": "{{ consideration_units_rt }}",
        "{{ consideration | reference_lines | nl2br }}": "{{ references_rt }}",
        "{{ consideration | reference_lines }}": "{{ references_rt }}",
    }

    block_loop_found = False
    inserted_block_loop = False

    def patch_paragraphs(paragraphs):
        nonlocal inserted_block_loop, block_loop_found
        for p in paragraphs:
            text = p.text
            if not text:
                continue

            if "{% for block in exp.blocks" in text:
                block_loop_found = True

            if "{% if block.type" in text and not block_loop_found and not inserted_block_loop:
                p.text = "{% for block in exp.blocks %}" + text
                inserted_block_loop = True
                block_loop_found = True
                text = p.text

            original = text
            modified = original
            for old, new in replacements.items():
                if old in modified:
                    modified = modified.replace(old, new)
            if modified != original:
                p.text = modified

    patch_paragraphs(docx_obj.paragraphs)
    for t in docx_obj.tables:
        for row in t.rows:
            for cell in row.cells:
                patch_paragraphs(cell.paragraphs)


def _prepare_template_docx_bytes(template_path: Path) -> bytes:
    """
    Pre-process the DOCX template *before* docxtpl parses it.

    Why: docxtpl/Jinja parsing happens against the raw XML from the template file. If Jinja tags
    are split across Word runs, docxtpl won't recognize them and will output the raw template text.
    """

    raw = template_path.read_bytes()
    doc = Document(BytesIO(raw))
    _normalize_jinja_tags(doc)

    # Apply our template patching on the raw Document so docxtpl sees the fixed tags during parsing.
    # (No need for `context` here; we only rewrite tag strings / insert missing loops.)
    replacements = {
        "{{ consideration.units | consideration_units | nl2br }}": "{{ consideration_units_rt }}",
        "{{ consideration.units | consideration_units }}": "{{ consideration_units_rt }}",
        "{{ consideration | reference_lines | nl2br }}": "{{ references_rt }}",
        "{{ consideration | reference_lines }}": "{{ references_rt }}",
    }

    block_loop_found = False
    inserted_block_loop = False

    def patch_paragraphs(paragraphs):
        nonlocal inserted_block_loop, block_loop_found
        for p in paragraphs:
            text = p.text
            if not text:
                continue

            if "{% for block in exp.blocks" in text:
                block_loop_found = True

            if "{% if block.type" in text and not block_loop_found and not inserted_block_loop:
                p.text = "{% for block in exp.blocks %}" + text
                inserted_block_loop = True
                block_loop_found = True
                text = p.text

            original = text
            modified = original
            for old, new in replacements.items():
                if old in modified:
                    modified = modified.replace(old, new)
            if modified != original:
                p.text = modified

    patch_paragraphs(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                patch_paragraphs(cell.paragraphs)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _normalize_jinja_tags(docx_obj) -> None:
    """
    DocxTemplate (docxtpl) requires Jinja tags like `{{ var }}` / `{% ... %}` to be contiguous.
    Word often splits them across multiple runs, which results in variables not being rendered.

    This function best-effort merges only the run fragments needed to make Jinja tags contiguous,
    without flattening whole paragraphs (to avoid losing formatting).
    """

    def normalize_paragraph(p) -> None:
        try:
            runs = list(p.runs)
        except Exception:
            return
        if not runs:
            return

        # Fast path: if nothing looks like Jinja, skip.
        try:
            joined = "".join((r.text or "") for r in runs)
        except Exception:
            return
        if "{{" not in joined and "{%" not in joined and "}}" not in joined and "%}" not in joined:
            return

        def merge_boundary(i: int) -> bool:
            if i + 1 >= len(runs):
                return False
            a = runs[i].text or ""
            b = runs[i + 1].text or ""
            if not a or not b:
                return False
            last = a[-1]
            first = b[0]
            # Stitch common split patterns: "{", "{", "{", "%", "}", "}", "%", "}".
            if last == "{" and first in {"{", "%"}:
                runs[i].text = a + b
                runs[i + 1].text = ""
                return True
            if last == "}" and first == "}":
                runs[i].text = a + b
                runs[i + 1].text = ""
                return True
            if last == "%" and first == "}":
                runs[i].text = a + b
                runs[i + 1].text = ""
                return True
            return False

        def find_start(text: str) -> tuple[int, str, str] | None:
            candidates: list[tuple[int, str, str]] = []
            i1 = text.find("{{")
            if i1 != -1:
                candidates.append((i1, "{{", "}}"))
            i2 = text.find("{%")
            if i2 != -1:
                candidates.append((i2, "{%", "%}"))
            if not candidates:
                return None
            return min(candidates, key=lambda t: t[0])

        i = 0
        while i < len(runs):
            if merge_boundary(i):
                continue

            text = runs[i].text or ""
            start = find_start(text)
            if not start:
                i += 1
                continue

            start_idx, _start_marker, end_marker = start
            # If end marker already exists in this run after the start, it's contiguous enough.
            if text.find(end_marker, start_idx + 2) != -1:
                i += 1
                continue

            # Otherwise, merge subsequent runs until we find the end marker.
            j = i
            while j < len(runs):
                if merge_boundary(j):
                    continue
                cur = runs[j].text or ""
                pos = cur.find(end_marker, (start_idx + 2) if j == i else 0)
                if pos == -1:
                    j += 1
                    continue

                # Found end marker in run j. Merge i..j into run i,
                # preserving any suffix after the end marker in run j.
                if j == i:
                    break

                prefix_end = cur[: pos + len(end_marker)]
                suffix = cur[pos + len(end_marker) :]

                merged = (runs[i].text or "") + "".join((runs[k].text or "") for k in range(i + 1, j)) + prefix_end
                runs[i].text = merged
                for k in range(i + 1, j):
                    runs[k].text = ""
                runs[j].text = suffix
                break

            i += 1

    try:
        for p in docx_obj.paragraphs:
            normalize_paragraph(p)
        for t in docx_obj.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        normalize_paragraph(p)
    except Exception:
        return

def _inject_inline_images(
    doc: DocxTemplate,
    context: dict,
    images_by_id: dict[str, ImageAsset],
    storage: Storage,
    *,
    image_bytes_by_id: dict[str, bytes] | None = None,
) -> None:
    experiments = context.get("experiments") or []
    for exp in experiments:
        blocks = exp.get("blocks") or []
        for block in blocks:
            if not isinstance(block, dict) or block.get("type") != "figure":
                continue
            fig = block.get("figure")
            if not isinstance(fig, dict):
                continue
            image_id = fig.get("figure_image_id")
            if not image_id:
                continue
            raw: bytes | None = None
            if image_bytes_by_id is not None:
                raw = image_bytes_by_id.get(str(image_id))
            if raw is None:
                asset = images_by_id.get(str(image_id))
                if not asset:
                    continue
                raw = storage.get_bytes(asset.storage_key)
            fig["figure_image"] = InlineImage(
                doc,
                BytesIO(raw),
                width=Mm(TARGET_WIDTH_MM),
                height=Mm(TARGET_HEIGHT_MM),
            )


def _is_missing_cell_value(value: str) -> bool:
    s = (value or "").strip()
    return not s or s == _MISSING_CELL_MARK


def _apply_missing_cell_diagonal(cell) -> None:
    # Draw a top-right to bottom-left diagonal in empty cells.
    try:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        diag = borders.find(qn("w:tr2bl"))
        if diag is None:
            diag = OxmlElement("w:tr2bl")
            borders.append(diag)
        diag.set(qn("w:val"), "single")
        diag.set(qn("w:sz"), "4")
        diag.set(qn("w:space"), "0")
        diag.set(qn("w:color"), "auto")
    except Exception:
        return


def _build_table_subdoc(doc: DocxTemplate, rows: Any):
    if not isinstance(rows, list) or not rows:
        return None
    max_cols = max((len(r) for r in rows if isinstance(r, list)), default=0)
    if max_cols <= 0:
        return None
    sub = doc.new_subdoc()
    table = sub.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        row_list = row if isinstance(row, list) else []
        for c_idx in range(max_cols):
            val = ""
            if c_idx < len(row_list):
                val = "" if row_list[c_idx] is None else str(row_list[c_idx])
            cell = table.cell(r_idx, c_idx)
            if r_idx > 0 and _is_missing_cell_value(val):
                cell.text = ""
                _apply_missing_cell_diagonal(cell)
            else:
                cell.text = val
    return sub


def _inject_tables(doc: DocxTemplate, context: dict) -> None:
    experiments = context.get("experiments") or []
    for exp in experiments:
        blocks = exp.get("blocks") or []
        for block in blocks:
            if not isinstance(block, dict) or block.get("type") != "table":
                continue
            tbl = block.get("table")
            if not isinstance(tbl, dict):
                continue
            subdoc = _build_table_subdoc(doc, tbl.get("rows"))
            if subdoc is not None:
                tbl["body"] = subdoc


def _paragraph_text(p_elm) -> str:
    try:
        texts: list[str] = []
        for node in p_elm.iter():
            if node.tag == qn("w:t"):
                texts.append(node.text or "")
        return "".join(texts).strip()
    except Exception:
        return ""


def _set_keep_next(p_elm) -> None:
    """
    Word pagination controls:
    - keepNext: keep this paragraph with the next element
    - keepLines: keep lines together within this paragraph
    """
    try:
        p_pr = p_elm.get_or_add_pPr()
        if p_pr.find(qn("w:keepNext")) is None:
            p_pr.append(OxmlElement("w:keepNext"))
        if p_pr.find(qn("w:keepLines")) is None:
            p_pr.append(OxmlElement("w:keepLines"))
    except Exception:
        return


def _set_row_cant_split(row) -> None:
    """
    Prevent a *single row* from splitting across pages (Word's "Allow row to break across pages" = off).
    """
    try:
        tr = row._tr  # python-docx internal
        tr_pr = tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
    except Exception:
        return


def _looks_like_table_caption(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    low = t.lower()
    if "表" in t or "table" in low:
        return True
    # Short "caption-ish" line (e.g., "1-1: xxxx", "実験結果:", etc.)
    if len(t) <= 60 and (":" in t or "：" in t):
        return True
    return False


def _apply_table_pagination_rules(docx_obj) -> None:
    """
    Best-effort Word pagination controls:
    - Keep table captions together with the following table.
    - Avoid splitting table rows across pages.
    - Encourage keeping the whole table together when it fits.

    Notes:
    - Word cannot *guarantee* an entire table stays on one page if it is taller than a page.
    """
    # 1) Keep captions with the table:
    #    - Always keep the nearest non-empty paragraph with the following table.
    #    - Also keep 1 more preceding paragraph when it looks like a caption line.
    try:
        body = docx_obj.element.body
        children = list(body.iterchildren())
        for i, child in enumerate(children):
            if child.tag != qn("w:tbl"):
                continue
            candidates: list[Any] = []
            j = i - 1
            while j >= 0 and len(candidates) < 2:
                prev = children[j]
                j -= 1
                if prev.tag != qn("w:p"):
                    break
                text = _paragraph_text(prev)
                if not text:
                    continue
                if not candidates:
                    candidates.append(prev)
                    continue
                if _looks_like_table_caption(text):
                    candidates.append(prev)
                break

            # candidates[0] is nearest to the table; candidates[1] keeps with candidates[0].
            for p in candidates:
                _set_keep_next(p)
    except Exception:
        pass

    # 2) Avoid splitting rows across pages & encourage keeping the whole table together.
    try:
        for table in docx_obj.tables:
            rows = table.rows
            for r_idx, row in enumerate(rows):
                _set_row_cant_split(row)
                if r_idx >= len(rows) - 1:
                    continue
                for cell in row.cells:
                    for p in cell.paragraphs:
                        try:
                            p.paragraph_format.keep_with_next = True
                            p.paragraph_format.keep_together = True
                        except Exception:
                            continue
    except Exception:
        return


def render_docx_bytes(
    *,
    template_path: str,
    context: dict,
    storage: Storage,
    job_id: str,
    image_bytes_by_id: dict[str, bytes] | None = None,
) -> bytes:
    template_file = Path(template_path)
    if not template_file.exists():
        raise FileNotFoundError(f"Template not found: {template_file}")

    prepared_bytes = _prepare_template_docx_bytes(template_file)
    tmp_path: str | None = None
    try:
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(prepared_bytes)
            tmp.flush()
            tmp_path = tmp.name

        if os.environ.get("DOCX_DEBUG") in {"1", "true", "True"}:
            try:
                storage.put_bytes(f"jobs/{job_id}/debug/template_prepared.docx", prepared_bytes)
            except Exception:
                pass

        doc = DocxTemplate(tmp_path)

        # Pre-calculate rich text values and patch template tags.
        context["consideration_units_rt"] = _create_consideration_units_rt(context.get("consideration", {}).get("units"))
        context["references_rt"] = _create_references_rt(context.get("consideration", {}))
        _patch_template(doc, context)

        images_by_id: dict[str, ImageAsset] = {}

        # We encode assets into the context when we call render (preferred).
        assets_images = context.pop("__assets_images", None)
        if isinstance(assets_images, dict):
            images_by_id = {k: ImageAsset.model_validate(v) for k, v in assets_images.items()}

        _inject_inline_images(doc, context, images_by_id, storage, image_bytes_by_id=image_bytes_by_id)
        _inject_tables(doc, context)

        env = _build_env()
        try:
            doc.render(context, jinja_env=env)
        except TemplateSyntaxError as exc:
            raise RuntimeError(f"Docx template parse failed: {exc}") from exc

        out = BytesIO()
        doc.save(out)
        rendered = out.getvalue()

        # Post-process the rendered document (python-docx) for pagination rules.
        # Applying python-docx mutations on DocxTemplate before save can cause docxtpl to
        # fall back to saving the original template (leaving tags unrendered).
        try:
            post = Document(BytesIO(rendered))
            _apply_table_pagination_rules(post)
            out2 = BytesIO()
            post.save(out2)
            return out2.getvalue()
        except Exception:
            return rendered
    finally:
        if tmp_path:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                pass
