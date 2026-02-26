from __future__ import annotations

import argparse
import json
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.config import load_settings
from core.storage import build_storage
from models.contracts import ImageAsset


def _load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _chapter_number(data: dict) -> str:
    chapter = str(((data.get("b_layer_bundle") or {}).get("method") or {}).get("chapter") or "").strip()
    if not chapter:
        return "?"
    try:
        return str(int(chapter) + 1)
    except ValueError:
        return chapter


def _parent_key(entry: dict) -> str:
    parent = entry.get("experiment_parent") or {}
    exp_key = str(parent.get("exp_key") or "").strip()
    if not exp_key:
        return ""
    parts = exp_key.split(".")
    if not parts:
        return exp_key
    try:
        parts[0] = str(int(parts[0]) + 1)
        return ".".join(parts)
    except ValueError:
        return exp_key


def _parent_title(entry: dict) -> str:
    parent = entry.get("experiment_parent") or {}
    return str(parent.get("title") or "").strip()


def _exp_key(entry: dict) -> str:
    return str(entry.get("experiment_number") or entry.get("experiment_name") or entry.get("exp_key") or "").strip()


def _exp_title(entry: dict) -> str:
    return str(entry.get("experiment_name") or "").strip()


def _format_tables(tables: list[dict]) -> list[str]:
    blocks: list[str] = []
    for t in tables:
        label = str(t.get("label") or "").strip()
        caption = str(t.get("caption") or "").strip()
        header = " ".join([p for p in [label, caption] if p]).strip()
        if header:
            blocks.append(f"<div style=\"text-align: center;\">{header}</div>")
        rows = t.get("rows") or []
        if not rows:
            continue
        header_row = rows[0]
        blocks.append("| " + " | ".join(header_row) + " |")
        blocks.append("| " + " | ".join(["---"] * len(header_row)) + " |")
        for r in rows[1:]:
            blocks.append("| " + " | ".join([str(x) for x in r]) + " |")
        blocks.append("")
    return blocks


def _format_graphs(graphs: list[dict]) -> list[str]:
    blocks: list[str] = []
    for g in graphs:
        label = str(g.get("label") or "").strip()
        caption = str(g.get("caption") or "").strip()
        image_id = str(g.get("image_id") or "").strip()
        if image_id:
            blocks.append(f"![](image:{image_id})")
        footer = " ".join([p for p in [label, caption] if p]).strip()
        if footer:
            blocks.append(f"<div style=\"text-align: center;\">{footer}</div>")
        blocks.append("")
    return blocks


def main() -> None:
    parser = argparse.ArgumentParser(description="Render L-layer docx from J-layer JSON.")
    parser.add_argument("--input", default="tmp_state_outputs/j_layer_merged.json")
    parser.add_argument("--out", default="tmp_state_outputs/l_layer_report.docx")
    parser.add_argument("--reference-doc", default="tmp_state_outputs/reference.docx")
    args = parser.parse_args()

    input_path = Path(args.input)
    out_path = Path(args.out)
    if not input_path.exists():
        raise SystemExit(f"INPUT_NOT_FOUND: {input_path}")

    data = _load_json(input_path)
    chapter = _chapter_number(data)
    pages = list(data.get("result_page") or [])

    lines: list[str] = []
    lines.append(f"# {chapter}.実験結果")
    lines.append("")

    by_parent: dict[str, list[dict]] = {}
    parent_titles: dict[str, str] = {}
    for page in pages:
        pkey = _parent_key(page)
        by_parent.setdefault(pkey, []).append(page)
        if pkey and pkey not in parent_titles:
            parent_titles[pkey] = _parent_title(page)

    for parent_key in sorted(by_parent.keys(), key=lambda k: k or ""):
        parent_title = parent_titles.get(parent_key, "")
        if parent_key:
            lines.append(f"## {parent_key} {parent_title}".strip())
            lines.append("")
        for page in by_parent[parent_key]:
            exp_key = _exp_key(page)
            exp_title = _exp_title(page)
            lines.append(f"### {exp_key} {exp_title}".strip())
            lines.append("")
            method_summary = str(page.get("method_summary") or "").strip()
            result_description = str(page.get("result_description") or "").strip()
            if method_summary:
                lines.append(method_summary)
                lines.append("")
            if result_description:
                lines.append(result_description)
                lines.append("")
            tables = list(page.get("tables") or [])
            graphs = list(page.get("graphs") or [])
            lines.extend(_format_tables(tables))
            lines.extend(_format_graphs(graphs))
            quant = page.get("quant_comment") or []
            if quant:
                text = str(quant[0].get("quant_comment") or "").strip()
                if text:
                    lines.append(text)
                    lines.append("")

    markdown = "\n".join(lines).strip() + "\n"
    settings = load_settings()
    storage = build_storage(backend=settings.storage_backend, storage_dir=settings.storage_dir)
    images_by_id = {img["image_id"]: ImageAsset.model_validate(img) for img in data.get("assets_images") or [] if img.get("image_id")}

    with tempfile.TemporaryDirectory(prefix="reportlab_pandoc_") as tmpdir:
        tmpdir_path = Path(tmpdir)
        image_map: dict[str, Path] = {}
        for image_id, asset in images_by_id.items():
            storage_key = str(asset.storage_key or "").strip()
            if not storage_key:
                continue
            try:
                raw = storage.get_bytes(storage_key)
            except Exception:
                continue
            suffix = Path(str(asset.filename or "")).suffix or ".png"
            out_img = tmpdir_path / f"{image_id}{suffix}"
            out_img.write_bytes(raw)
            image_map[image_id] = out_img

        md_path = tmpdir_path / "input.md"
        for image_id, img_path in image_map.items():
            markdown = markdown.replace(f"![](image:{image_id})", f"![]({img_path})")
        md_path.write_text(markdown, encoding="utf-8")

        cmd = ["pandoc", str(md_path), "-o", str(out_path)]
        ref_path = Path(args.reference_doc)
        if ref_path.exists():
            cmd.extend(["--reference-doc", str(ref_path)])
        subprocess.run(cmd, check=True)

    # Post-process: ensure table grid lines and caption alignment.
    try:
        doc = Document(out_path)
        caption_re = re.compile(r"^(図|表)\\s*\\d+(?:\\.\\d+){0,3}(?:\\.\\d+)?\\s*[:：]?\\s*.+")
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text and caption_re.match(text):
                try:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
        for table in doc.tables:
            try:
                table.style = "Table Grid"
            except Exception:
                pass
        doc.save(out_path)
    except Exception:
        pass

    print(out_path)


if __name__ == "__main__":
    main()
