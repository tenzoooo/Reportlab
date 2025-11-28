from docx import Document
from docx.shared import Pt

def create_template(output_path):
    doc = Document()
    
    # Title
    doc.add_heading('{{ chapter_title }}', level=1)
    
    # Section Loop
    p = doc.add_paragraph('{% for section in sections %}')
    
    # Section Header
    doc.add_heading('{{ section.section_number }} {{ section.title }}', level=2)
    
    # Subsection Loop
    doc.add_paragraph('{% for subsection in section.subsections %}')
    
    # Subsection Header
    doc.add_heading('{{ subsection.subsection_number }} {{ subsection.title }}', level=3)
    
    # Content Block Loop
    doc.add_paragraph('{% for block in subsection.content_blocks %}')
    
    # Text Block
    doc.add_paragraph('{% if block.type == "text" %}{{ block.content }}{% endif %}')
    
    # Table Block (Placeholder - actual table rendering needs subdoc)
    # We use a special tag that our renderer will replace with a subdoc
    doc.add_paragraph('{% if block.type == "table" %}{{ block.content }}{% endif %}')
    
    # Figure Block
    doc.add_paragraph('{% if block.type == "figure" %}{{ block.content }}{% endif %}')
    
    doc.add_paragraph('{% endfor %}') # End Content Block Loop
    doc.add_paragraph('{% endfor %}') # End Subsection Loop
    doc.add_paragraph('{% endfor %}') # End Content Block Loop
    doc.add_paragraph('{% endfor %}') # End Subsection Loop
    doc.add_paragraph('{% endfor %}') # End Section Loop

    # --- Consideration Section ---
    doc.add_page_break()
    doc.add_heading('2. 考察', level=1)
    
    # Using the custom filter 'consideration_units' which formats the units list
    # The renderer maps 'consideration_units_rt' to the context
    doc.add_paragraph('{{ consideration_units_rt }}')

    # --- Summary Section ---
    doc.add_page_break()
    doc.add_heading('3. まとめ', level=1)
    doc.add_paragraph('{{ summary }}')

    # --- References Section ---
    doc.add_page_break()
    doc.add_heading('参考文献', level=1)
    
    # Using the custom filter 'reference_lines' which formats the references
    # The renderer maps 'references_rt' to the context
    doc.add_paragraph('{{ references_rt }}')

    
    doc.save(output_path)
    print(f"Created template at {output_path}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python3 create_template.py <output_path>")
        sys.exit(1)
    create_template(sys.argv[1])
