
import sys
import os

# Add the directory containing the library to the python path
sys.path.append(os.path.join(os.getcwd(), 'lib/docx'))

from docxtpl import DocxTemplate
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from render_with_docxtpl import build_table_subdoc

def verify_alignment():
    # Create a dummy template (needs a valid docx file to init, can use any)
    # We'll create a blank one using python-docx first if needed, 
    # but DocxTemplate needs a file.
    # Let's assume we can use 'templates/template.docx' which is smaller.
    
    template_path = 'templates/template.docx'
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = DocxTemplate(template_path)
    
    rows = [
        ["Header 1", "Header 2"],
        ["Cell 1", "Cell 2"]
    ]
    
    subdoc = build_table_subdoc(doc, rows)
    
    # Access the table in the subdoc
    # subdoc is a docx.document.Document object (or similar proxy)
    # We need to find the table we just added.
    
    if not subdoc.tables:
        print("No tables found in subdoc")
        sys.exit(1)
        
    table = subdoc.tables[0]
    
    # Check alignment of the first cell
    cell = table.cell(0, 0)
    
    print(f"Vertical Alignment: {cell.vertical_alignment}")
    if cell.vertical_alignment != WD_ALIGN_VERTICAL.CENTER:
        print("FAIL: Vertical alignment is not CENTER")
        sys.exit(1)
        
    # Check paragraph alignment
    if not cell.paragraphs:
        print("FAIL: No paragraphs in cell")
        sys.exit(1)
        
    para_alignment = cell.paragraphs[0].alignment
    print(f"Paragraph Alignment: {para_alignment}")
    
    if para_alignment != WD_ALIGN_PARAGRAPH.CENTER:
        print("FAIL: Paragraph alignment is not CENTER")
        sys.exit(1)

    print("SUCCESS: Table cells are center aligned.")

if __name__ == "__main__":
    verify_alignment()
