from fastapi import FastAPI, Response
from pydantic import BaseModel, Field
from typing import List, Optional, Any
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import base64, io, os

TEMPLATE_PATH_DEFAULT = r"/mnt/data/実験結果_テンプレ_v3_centered.docx"

app = FastAPI()

class TableSpec(BaseModel):
    label: str
    caption: str
    header: Optional[List[str]] = None
    rows: Optional[List[List[Any]]] = None
    # if provided, overrides header/rows and is used as a raw subdoc placeholder is not possible via API
    # We will always build subdoc from header/rows for simplicity.

class FigureSpec(BaseModel):
    label: str
    caption: str
    image_b64: Optional[str] = None
    image_path: Optional[str] = None
    width_mm: Optional[float] = 100.0

class Experiment(BaseModel):
    idx: int
    subidx: Optional[int] = None
    name: str
    description_brief: str
    tables: Optional[List[TableSpec]] = None
    figures: Optional[List[FigureSpec]] = None
    quant_comment: str

class Payload(BaseModel):
    chapter: int
    experiments: List[Experiment]
    # Optional template data: base64 or filesystem path
    template_b64: Optional[str] = None
    template_path: Optional[str] = None

@app.get("/health")
def health():
    return {"status": "ok"}

def load_template(payload: Payload) -> DocxTemplate:
    if payload.template_b64:
        tpl_bytes = base64.b64decode(payload.template_b64)
        bio = io.BytesIO(tpl_bytes)
        return DocxTemplate(bio)
    path = payload.template_path or TEMPLATE_PATH_DEFAULT
    if not os.path.exists(path):
        raise FileNotFoundError(f"Template not found: {path}")
    return DocxTemplate(path)

def build_table_subdoc(tpl: DocxTemplate, header: List[str], rows: List[List[Any]]):
    sd = tpl.new_subdoc()
    # create table
    t = sd.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    # header
    for i, h in enumerate(header):
        t.cell(0, i).text = str(h)
    # rows
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = "" if v is None else str(v)
    return sd

@app.post("/render")
def render_doc(payload: Payload):
    tpl = load_template(payload)

    # Build context compatible with the v3 template:
    # - tables use 'doc' (subdoc) placeholder
    # - figures use InlineImage
    exp_ctx = []
    for exp in payload.experiments:
        e = {
            "idx": exp.idx,
            "subidx": exp.subidx,
            "name": exp.name,
            "description_brief": exp.description_brief,
            "quant_comment": exp.quant_comment,
        }
        # tables
        tbls = []
        if exp.tables:
            for t in exp.tables:
                header = t.header or []
                rows = t.rows or []
                sd = build_table_subdoc(tpl, header, rows) if header else tpl.new_subdoc()
                tbls.append({
                    "label": t.label,
                    "caption": t.caption,
                    "doc": sd,
                })
        e["tables"] = tbls

        # figures
        figs = []
        if exp.figures:
            for f in exp.figures:
                img = None
                if f.image_b64:
                    img_bytes = base64.b64decode(f.image_b64)
                    img = InlineImage(tpl, io.BytesIO(img_bytes), width=Mm(f.width_mm or 100.0))
                elif f.image_path:
                    img = InlineImage(tpl, f.image_path, width=Mm(f.width_mm or 100.0))
                figs.append({
                    "label": f.label,
                    "caption": f.caption,
                    "image": img
                })
        e["figures"] = figs
        exp_ctx.append(e)

    context = {
        "chapter": payload.chapter,
        "experiments": exp_ctx
    }

    tpl.render(context)
    out = io.BytesIO()
    tpl.save(out)
    data = out.getvalue()
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=experiment_results.docx"}
    )
